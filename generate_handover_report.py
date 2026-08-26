import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import os

def build_handover_report():
    doc = docx.Document()

    # Page setup - Standard A4, 1-inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.page_width = Inches(8.27)  # A4
        section.page_height = Inches(11.69)
        
        # Header & Footer setup
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("TVET Agri – Ubuhinzi Skills+ Project | TVET Expert Handover Report")
        hrun.font.name = 'Calibri'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(120, 120, 120)
        
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        frun1 = fp.add_run("Expertise France / European Union (Contract 700002154)   |   ")
        frun1.font.name = 'Calibri'
        frun1.font.size = Pt(8.5)
        frun1.font.color.rgb = RGBColor(120, 120, 120)
        frun2 = fp.add_run("Page ")
        frun2.font.name = 'Calibri'
        frun2.font.size = Pt(8.5)
        frun2.font.color.rgb = RGBColor(120, 120, 120)
        
        fldSimple = OxmlElement('w:fldSimple')
        fldSimple.set(qn('w:instr'), 'PAGE')
        fp._p.append(fldSimple)

    # Styles Setup
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(40, 40, 40)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(4)

    # Helper Functions
    def set_cell_background(cell, hex_color):
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
        cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

    def set_cell_margins(cell, top=100, bottom=100, left=130, right=130):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{m}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def set_table_borders(table, color="D0D7DE", sz="4", val="single"):
        tblPr = table._tbl.tblPr
        borders_xml = f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
        '''
        tblPr.append(parse_xml(borders_xml))

    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(27, 54, 93) # Navy Blue
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(46, 117, 182) # Steel Blue
        return h

    def add_heading_3(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = RGBColor(70, 70, 70) # Charcoal
        return h

    def add_paragraph(text, bold_prefix=None, space_after=4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Calibri'
            r_bold.font.size = Pt(9.5)
            r_bold.font.bold = True
            r_bold.font.color.rgb = RGBColor(30, 30, 30)
        if text:
            r_text = p.add_run(text)
            r_text.font.name = 'Calibri'
            r_text.font.size = Pt(9.5)
            r_text.font.color.rgb = RGBColor(50, 50, 50)
        return p

    def add_bullet(text, bold_prefix=None, level=0):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(2.5)
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Calibri'
            r_bold.font.size = Pt(9.5)
            r_bold.font.bold = True
            r_bold.font.color.rgb = RGBColor(30, 30, 30)
        if text:
            r_text = p.add_run(text)
            r_text.font.name = 'Calibri'
            r_text.font.size = Pt(9.5)
            r_text.font.color.rgb = RGBColor(50, 50, 50)
        return p

    def add_callout(text, title=None, alert_type="note"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        cell.width = Inches(6.27)
        
        if alert_type == "warning":
            bg_color = "FFFBEB"
            border_color = "D97706"
            title_color = RGBColor(180, 83, 9)
        elif alert_type == "important":
            bg_color = "FEF2F2"
            border_color = "DC2626"
            title_color = RGBColor(185, 28, 28)
        elif alert_type == "success":
            bg_color = "F0FDF4"
            border_color = "16A34A"
            title_color = RGBColor(21, 128, 61)
        else:
            bg_color = "F0F4F8"
            border_color = "2E75B6"
            title_color = RGBColor(31, 78, 121)
            
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
        
        tcPr = cell._tc.get_or_add_tcPr()
        borders_xml = f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="none"/>
            <w:bottom w:val="none"/>
            <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>
            <w:right w:val="none"/>
        </w:tcBorders>
        '''
        tcPr.append(parse_xml(borders_xml))
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.15
        if title:
            r_title = p.add_run(f"[{title.upper()}] ")
            r_title.font.name = 'Calibri'
            r_title.font.size = Pt(9.5)
            r_title.font.bold = True
            r_title.font.color.rgb = title_color
        
        r_body = p.add_run(text)
        r_body.font.name = 'Calibri'
        r_body.font.size = Pt(9.5)
        r_body.font.color.rgb = RGBColor(40, 40, 40)
        
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(3)

    def add_styled_table(headers, rows_data, col_widths=None, alignment=None):
        tbl = doc.add_table(rows=len(rows_data) + 1, cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(tbl, color="D0D7DE", sz="4", val="single")
        
        hdr_cells = tbl.rows[0].cells
        for i, title in enumerate(headers):
            hdr_cells[i].text = title
            set_cell_background(hdr_cells[i], "1B365D")
            set_cell_margins(hdr_cells[i], top=80, bottom=80, left=80, right=80)
            p = hdr_cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            if alignment and i < len(alignment):
                p.alignment = alignment[i]
            for run in p.runs:
                run.font.name = 'Calibri'
                run.font.size = Pt(8.5)
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                
        trPr = tbl.rows[0]._tr.get_or_add_trPr()
        trPr.append(parse_xml(f'<w:tblHeader {nsdecls("w")}/>'))

        for r_idx, row in enumerate(rows_data):
            row_cells = tbl.rows[r_idx + 1].cells
            r_trPr = tbl.rows[r_idx + 1]._tr.get_or_add_trPr()
            r_trPr.append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))
            
            bg_col = "FFFFFF" if r_idx % 2 == 0 else "F8FAFC"
            for c_idx, val in enumerate(row):
                row_cells[c_idx].text = str(val) if val is not None else ""
                set_cell_background(row_cells[c_idx], bg_col)
                set_cell_margins(row_cells[c_idx], top=60, bottom=60, left=70, right=70)
                p = row_cells[c_idx].paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.1
                if alignment and c_idx < len(alignment):
                    p.alignment = alignment[c_idx]
                for run in p.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(8)
                    run.font.color.rgb = RGBColor(50, 50, 50)
                    
                    t = run.text.strip()
                    if t == "COMPLETED":
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(22, 101, 52)
                    elif t == "IN PROGRESS":
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(30, 64, 175)
                    elif t in ["DELAYED / AT RISK", "AT RISK", "DELAYED"]:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(185, 28, 28)
                    elif t == "PLANNED":
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(107, 114, 128)
                    elif t == "HIGH":
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(185, 28, 28)
                    elif t == "MEDIUM":
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(217, 119, 6)
                    elif t == "LOW":
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(107, 114, 128)

        if col_widths:
            for r in tbl.rows:
                for c_idx, w in enumerate(col_widths):
                    r.cells[c_idx].width = Inches(w)

        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(3)
        return tbl

    # ==========================================
    # DOCUMENT HEADER / TITLE BLOCK
    # ==========================================
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run("TVET EXPERT HANDOVER REPORT")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(20)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(27, 54, 93)

    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_after = Pt(10)
    r_sub = p_sub.add_run("TVET Agri – Ubuhinzi Skills+ Project  |  Expertise France Rwanda")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(12)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(46, 117, 182)

    # Metadata Info Table
    meta_headers = ["Project Parameter", "Details & Administrative Context"]
    meta_data = [
        ["Project Title", "Agricultural Transformation Through Skills Upgrading of Women and Youth in Rwanda (TVET Agri – Ubuhinzi Skills+)"],
        ["Contract & Donor", "EU Contract Number: 700002154 | European Union / Team Europe Initiative (TEI)"],
        ["Implementing Agency", "Expertise France (EF), Rwanda Country Office"],
        ["Role / Assignment", "Technical Vocational Education and Training (TVET) Expert"],
        ["Target Audience", "Incoming TVET Expert (Successor) & Project Team Leader"],
        ["Reporting Period / Baseline", "Assignment Handover (Baseline Activity Matrix as of 15 July 2026 / Local Handover: August 2026)"],
        ["Core Project Focus", "Output 1.2 (Standardised Quality Agri-TVET Training) & Output 1.3 (Agri CoVEs/CoEs Capacity & Infrastructure)"]
    ]
    add_styled_table(meta_headers, meta_data, col_widths=[2.0, 4.27])

    # ==========================================
    # 1. EXECUTIVE SUMMARY
    # ==========================================
    add_heading_1("1. Executive Summary")
    
    add_paragraph("This Handover Report provides a concise, professional, evidence-based, and action-oriented synthesis of the technical portfolio managed by the outgoing TVET Expert under the EU-funded ", bold_prefix="Assignment Purpose: ")
    p_exec = doc.paragraphs[-1]
    r_ex1 = p_exec.add_run("TVET Agri – Ubuhinzi Skills+ Project (2024–2029)")
    r_ex1.font.bold = True
    r_ex2 = p_exec.add_run(", implemented by Expertise France in Rwanda. The primary objective is to ensure a seamless operational transition for the incoming TVET Expert and provide the Project Team Leader with an unambiguous status map of all deliverables, active consultancies, strategic dependencies, critical bottlenecks, and required decisions.")

    add_paragraph("The TVET Expert oversees the technical core of Outcome 1, leading ", bold_prefix="Scope of Portfolio: ")
    p_sc = doc.paragraphs[-1]
    p_sc.add_run("Output 1.2 (Enhanced capacity of Agri TVET Centres at national level to provide standardized quality training relevant to labour market needs, climate action, and gender responsiveness)")
    p_sc.add_run(" and ")
    p_sc.add_run("Output 1.3 (Enhanced capacity of targeted Agri Centres of Vocational Excellence - CoVEs/CoEs - to provide quality training and a conducive learning environment).")

    add_heading_2("Executive Snapshot of Portfolio Status")
    add_bullet("Completed Foundational Assets: Comprehensive technical review and structuring of the Practical Training Unit (PTU) Model Technical Design for TVET Centres of Excellence at Kisaro TSS (Rulindo District) and EFA Nyagahanga TSS (Gatsibo District); standardisation of the project-wide Capacity Building Concept Note Framework; and baseline structuring of the 2026 TVET Activity Matrix.", bold_prefix="Accomplished: ")
    add_bullet("Active In-Flight Workstreams: 4 major activity packages (7 individual sub-activities) are actively in progress under Output 1.2, covering (i) Experiential Learning Assessment Guidelines, (ii) Digital CBT/CBA Learning Content across Agriculture, Animal Health, and Food Processing, (iii) Gender and Social Inclusion Curricula Review, and (iv) Gender-Responsive Trainer's Guide.", bold_prefix="In Progress: ")
    add_bullet("Critical Bottlenecks & Strategic Risks: (i) Activity A.1.2.6 (Modular Short Courses) is delayed due to an external dependency on the overarching Needs Assessment led by Aimable; (ii) Activity A.1.2.2 (Agri CoE Accreditation Framework) carries high duplication risk with parallel initiatives led by ETF/RTB and must be clarified before proceeding; (iii) Activity A.1.3.8 (Resource Centre) has been rescheduled to 2028.", bold_prefix="Delayed / At Risk: ")
    add_bullet("Sequential Pipeline (2026–2028): Multi-year workpackages scheduled to start include CoE Governance Models (A.1.2.1; Sept 2026–2027), CBT/CBA Staff Capacity Building (A.1.2.8; Aug 2026–2027), Recognition of Prior Learning (RPL) in Agriculture (A.1.2.9; Oct 2026–2028), Excellency Committees (A.1.3.1; 2027), and Teacher Workplace Internships (A.1.3.6; 2027).", bold_prefix="Planned Pipeline: ")

    add_callout(
        "Immediate priority must be given to supervising the four active consultancy workstreams (Guidelines, Curricula, GESI Review, Trainer's Guide) through to their national validation workshops, while resolving the Needs Assessment dependency to unlock the procurement of 3 short-course consultants.",
        title="Key Management Takeaway",
        alert_type="important"
    )

    # ==========================================
    # 2. OVERALL STATUS
    # ==========================================
    add_heading_1("2. Overall Status & Portfolio Analytics")
    
    add_paragraph("The TVET Expert assignment encompasses 12 main activity packages across Outputs 1.2 and 1.3, comprising 26 granular sub-activity lines tracked within the project's 2026 Workplan. The portfolio represents a balance of foundational pedagogical reform, digital content creation, institutional governance, and school-level infrastructure enablement.")

    # Status Distribution Table
    stat_headers = ["Category / Status", "Main Activities", "Sub-Activities", "Operational Share (%)", "Key Focus Areas"]
    stat_data = [
        ["COMPLETED", "2 Foundational", "N/A", "10%", "PTU Technical Design Review (Kisaro & Nyagahanga), Capacity Building Template, Workplan Baseline"],
        ["IN PROGRESS", "4 Packages", "7 Lines", "30%", "Experiential Assessment Guidelines (A.1.2.3), Digital CBT/CBA Modules (A.1.2.4), Gender Review (A.1.2.5), Trainer's Guide (A.1.2.7)"],
        ["DELAYED / AT RISK", "3 Packages", "6 Lines", "25%", "Modular Short Courses (A.1.2.6 - Needs Assessment dependency); CoE Accreditation (A.1.2.2 - Feasibility/Overlap); Resource Centre (A.1.3.8 - Shifted to 2028)"],
        ["PLANNED", "7 Packages", "13 Lines", "35%", "Governance Models (A.1.2.1), CBT/CBA ToT (A.1.2.8), RPL TA (A.1.2.9), Excellency Committees (A.1.3.1), Teacher OJT (A.1.3.6)"],
        ["TOTAL PORTFOLIO", "12 Main Activities", "26 Tracked Lines", "100%", "Full Output 1.2 and Output 1.3 Delivery Framework"]
    ]
    add_styled_table(stat_headers, stat_data, col_widths=[1.5, 1.0, 0.9, 1.1, 1.77], alignment=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])

    add_heading_2("Project Results Framework Indicator Targets Managed by TVET Expert")
    add_paragraph("The TVET Expert directly contributes to several core indicators in the EU Description of Action (DoA) Results Framework:")
    add_bullet("Indicator 1.2.1: Number of guiding documents for sustainable Agri CoEs developed (Target: Governance model, PTU business model, trainer's guide).", bold_prefix="Guiding Documents: ")
    add_bullet("Indicator 1.2.2: Number of CBT/CBA curricula developed/reviewed in Agriculture, Animal Health, Food Processing (Target: 3 curricula, 2 digital module sets).", bold_prefix="Curricula & Modules: ")
    add_bullet("Indicator 1.2.3: Number of pedagogical staff and lead trainers trained in CBT/CBA and GEDI (Target: 4 capacity building sessions, 1 ToT for lead trainers, 1 ToT for trainers).", bold_prefix="Capacity Building: ")
    add_bullet("Indicator 1.2.4: Number of short courses developed on sustainable/circular economy, gender, and nutrition (Target: 2 annual modular sets).", bold_prefix="Short Courses: ")
    add_bullet("Indicator 1.2.5: Technical assistance to RTB for RPL in Agriculture (Target: 3 workshops, assessor training, system operationalization).", bold_prefix="RPL System: ")
    add_bullet("Indicator 1.3.1: Number of Agri CoVEs supported with functional PTUs, Excellency Committees, and Resource Centres (Target: 2 CoEs - Kisaro TSS and EFA Nyagahanga TSS).", bold_prefix="CoVE Support: ")

    # ==========================================
    # 3. KEY ACHIEVEMENTS
    # ==========================================
    add_heading_1("3. Key Achievements & Accomplished Deliverables")
    
    add_paragraph("During the assignment period, the TVET Expert established critical technical baselines and advanced core deliverables across curriculum development, institutional design, and capacity-building frameworks:")

    add_heading_2("3.1. Technical Review of Practical Training Unit (PTU) Model Design")
    add_paragraph("The TVET Expert conducted an extensive technical review and quality assurance of the comprehensive ", bold_prefix="Milestone Accomplishment: ")
    p_ptu = doc.paragraphs[-1]
    p_ptu.add_run("Technical Design for Practical Training Units (PTU) of TVET Centres of Excellence")
    p_ptu.add_run(" for ")
    p_ptu.add_run("Kisaro TSS (Rulindo District)")
    p_ptu.add_run(" and ")
    p_ptu.add_run("EFA Nyagahanga TSS (Gatsibo District)")
    p_ptu.add_run(". This 50+ page blueprint establishes a modern circular bio-economy model combining hands-on technical training, commercial production units, and technology demonstration.")
    add_bullet("Integrated Zoning & Spatial Allocation: Defined technical specifications, land requirements, and machinery for 6 specialized sub-units: (1) Crop Production & Horticulture, (2) Food Processing & Post-Harvest Value Addition, (3) Animal Health & Veterinary Clinic, (4) Climate-Smart Irrigation & Water Management, (5) Adult Learning & RPL Centre, and (6) Circular Waste-to-Energy (Biogas, Biofertilizer composting, Wastewater treatment, and Machinery Storage / 'Mini Garage').", bold_prefix="Circular Infrastructure: ")
    add_bullet("Dual Accountability Governance Model: Established the organizational framework balancing pedagogical objectives with commercial agribusiness viability, including unit manager responsibilities and revenue reinvestment mechanisms.", bold_prefix="Governance Structure: ")
    add_bullet("Quality Assurance & Standards Alignment: Embedded Standard Operating Procedures (SOPs) and safety compliance benchmarking against Rwanda Agriculture Board (RAB), Rwanda Food and Drugs Authority (FDA), and HACCP/GMP standards.", bold_prefix="QA & SOP Integration: ")

    add_heading_2("3.2. Curricula Modernization & Digital CBT/CBA Content Formulation")
    add_paragraph("Initiated and technically directed the drafting of digital learning modules and CBT/CBA training manuals across three flagship agricultural trades: ", bold_prefix="Pedagogical Reform: ")
    p_cur = doc.paragraphs[-1]
    p_cur.add_run("Crop Production (Agriculture), Animal Health (Veterinary), and Food Processing")
    p_cur.add_run(". The content integrates green competencies, climate-resilient practices, and interactive digital pedagogy.")

    add_heading_2("3.3. Mainstreaming Gender Equality, Disability & Social Inclusion (GEDI)")
    add_paragraph("Conducted technical screening and launched the dedicated curricula review to systematically embed gender-responsive and disability-inclusive training practices. Oversaw the drafting of the specialized ", bold_prefix="Inclusion Framework: ")
    p_gedi = doc.paragraphs[-1]
    p_gedi.add_run("Gender-Responsive and Inclusive Trainers' Guide (Activity A.1.2.7)")
    p_gedi.add_run(", providing TVET instructors with concrete pedagogical tools to address the specific needs of female students and trainees with disabilities.")

    add_heading_2("3.4. Pedagogical Capacity Building Framework Standardisation")
    add_paragraph("Developed and operationalized the standardized ", bold_prefix="Methodological Asset: ")
    p_cn = doc.paragraphs[-1]
    p_cn.add_run("Training / Capacity Building Concept Note Template")
    p_cn.add_run(". This framework standardizes all future project training events, ensuring strict alignment with Bloom's Taxonomy learning outcomes, the project results framework, M&E pre/post-testing protocols, and cost-efficiency benchmarks.")

    # ==========================================
    # 4. ONGOING & PENDING ACTIVITIES
    # ==========================================
    add_heading_1("4. Ongoing & Pending Activities (Active Technical Streams)")
    
    add_paragraph("Four primary activity packages under Output 1.2 are actively underway. The successor must maintain continuous oversight of consultant drafting schedules and coordinate upcoming multi-stakeholder validation workshops:")

    # Detailed table of active streams
    act_headers = ["Activity Ref & Title", "Target / Indicator", "Current Status", "Work Accomplished", "Pending Deliverables", "Next Action Required"]
    act_data = [
        [
            "A.1.2.3 Develop Assessment Guidelines for Agri CoEs",
            "1 Guideline Document",
            "IN PROGRESS",
            "Consultant contracted (1.2.3.1); Draft experiential learning assessment model developed (1.2.3.2).",
            "Draft guideline under internal technical review; validation workshop pending (1.2.3.3).",
            "Review draft guidelines; organize national validation workshop with RTB/NESA."
        ],
        [
            "A.1.2.4 Develop/Review Curricula in Agri, Animal Health, Food Processing",
            "3 Curricula; 2 Digital Module Sets",
            "IN PROGRESS",
            "Drafting of Digital Learning Content and CBT/CBA Training Manuals underway (1.2.4.1).",
            "Finalization of 2 digital module sets across 3 trades; validation workshops pending (1.2.4.2).",
            "Quality review of digital modules; convene 2 validation workshops with sector experts."
        ],
        [
            "A.1.2.5 Review Curricula for Gender-Responsiveness and Inclusion",
            "1 Workshop; 2 Digital Module Sets",
            "IN PROGRESS",
            "Curricula review underway; Drafting of GEDI digital modules in progress (1.2.5.1).",
            "Consolidated GEDI review report; module drafts; validation workshop pending (1.2.5.2).",
            "Collaborate with EF Gender Expert to validate modules; organize validation workshop."
        ],
        [
            "A.1.2.7 Develop Gender-Responsive Trainers' Guide",
            "1 Trainers' Guide",
            "IN PROGRESS",
            "Consultant contracted; Initial draft of Trainer's Guide developed (1.2.7.1).",
            "Technical refinement of Trainer's Guide; validation workshop pending (1.2.7.2).",
            "Complete review of draft guide; organize stakeholder validation workshop with lead trainers."
        ]
    ]
    add_styled_table(act_headers, act_data, col_widths=[1.3, 0.9, 0.9, 1.1, 1.0, 1.07])

    add_heading_2("Detailed Operational Breakdown of Active Workstreams")
    
    add_heading_3("Workstream 4.1: Assessment Guidelines & Experiential Learning Model (Activity A.1.2.3)")
    add_bullet("Target Deliverable: Validated Assessment Guidelines incorporating climate action, gender, and social inclusion criteria based on the experiential learning model.", bold_prefix="Objective: ")
    add_bullet("Current Status: Sub-activities 1.2.3.1 (Consultant contracting) and 1.2.3.2 (Guidelines development) are IN PROGRESS. Draft guidelines are undergoing technical consolidation.", bold_prefix="Progress: ")
    add_bullet("Immediate Next Steps: Provide consolidated technical feedback on the draft experiential learning model; draft concept note and budget for the validation workshop (Sub-activity 1.2.3.3 - PLANNED); coordinate participant invitations with RTB and NESA.", bold_prefix="Action: ")

    add_heading_3("Workstream 4.2: Curricula Modernization & Digital Learning Content (Activity A.1.2.4)")
    add_bullet("Target Deliverable: 3 Curricula reviewed (Agriculture, Animal Health, Food Processing) and 2 sets of interactive Digital Learning Content / CBT/CBA Training Manuals.", bold_prefix="Objective: ")
    add_bullet("Current Status: Sub-activity 1.2.4.1 (Digital content development) is IN PROGRESS. Authors are compiling modular digital packages.", bold_prefix="Progress: ")
    add_bullet("Immediate Next Steps: Conduct pedagogical review of draft modules against CBT/CBA requirements; ensure alignment with green competencies; schedule and organize two national validation workshops (Sub-activity 1.2.4.2 - PLANNED).", bold_prefix="Action: ")

    add_heading_3("Workstream 4.3: Gender-Responsive & Social Inclusion Curricula Review (Activity A.1.2.5)")
    add_bullet("Target Deliverable: Comprehensive review of curricula in 3 trades and 2 validated digital modules focusing on gender equality, disability inclusion, and anti-harassment.", bold_prefix="Objective: ")
    add_bullet("Current Status: Sub-activity 1.2.5.1 (Curricula review and module drafting) is IN PROGRESS in coordination with the Gender Expert.", bold_prefix="Progress: ")
    add_bullet("Immediate Next Steps: Finalize trade-specific GEDI module drafts; plan the stakeholder validation workshop (Sub-activity 1.2.5.2 - PLANNED) with RTB, GMO, and disability advocacy groups.", bold_prefix="Action: ")

    add_heading_3("Workstream 4.4: Gender-Responsive & Inclusive Trainers' Guide (Activity A.1.2.7)")
    add_bullet("Target Deliverable: 1 Practical Guide for TVET Instructors on delivering CBT/CBA training in a gender-responsive and disability-inclusive manner.", bold_prefix="Objective: ")
    add_bullet("Current Status: Sub-activity 1.2.7.1 is IN PROGRESS. The consultant has prepared the initial draft structure.", bold_prefix="Progress: ")
    add_bullet("Immediate Next Steps: Review draft guide for practical usability in classroom and farm environments; convene the validation workshop (Sub-activity 1.2.7.2 - PLANNED).", bold_prefix="Action: ")

    # ==========================================
    # 5. SCHOOL / WORKSTREAM STATUS
    # ==========================================
    add_heading_1("5. School & Workstream Status")
    
    add_paragraph("The TVET Expert assignment directly supports targeted technical secondary schools transitioning into Centres of Vocational Excellence (CoVEs / CoEs), alongside national-level thematic workstreams.")

    add_heading_2("5.1. Centre of Vocational Excellence (CoVE) School Sites")
    
    add_heading_3("Kisaro TSS (Rulindo District, Northern Province)")
    add_paragraph("Kisaro TSS is a primary investment site for infrastructure, commercial production, and pedagogical modernization:", bold_prefix="Site Profile: ")
    add_bullet("PTU Infrastructure Focus: Crop production sub-unit, modern horticulture greenhouses, climate-smart irrigation/water harvesting, animal health unit, biogas digester, and wastewater recycling.", bold_prefix="PTU Component: ")
    add_bullet("Governance & Management: Designated for tailored CoE governance model development (Activity 1.2.1.2) and future Excellency Committee establishment (Activity A.1.3.1).", bold_prefix="Governance: ")
    add_bullet("Current Needs / Outstanding: Finalize site validation for PTU construction; align school leadership with PTU commercialization and production unit operational modalities.", bold_prefix="Immediate Focus: ")

    add_heading_3("EFA Nyagahanga TSS (Gatsibo District, Eastern Province)")
    add_paragraph("EFA Nyagahanga TSS is designated as a flagship agricultural and animal husbandry CoE:", bold_prefix="Site Profile: ")
    add_bullet("PTU Infrastructure Focus: Food processing and value-addition unit (dairy/grain processing), animal health and veterinary demonstration sub-unit, adult learning/RPL facility, biofertilizer composting, and machinery mini-garage.", bold_prefix="PTU Component: ")
    add_bullet("Governance & Management: Designated for tailored CoE governance model development (Activity 1.2.1.2) and Excellency Action Plan rollout.", bold_prefix="Governance: ")
    add_bullet("Current Needs / Outstanding: Site-specific adaptation of the PTU architectural plans; formal agreement on resource centre space allocation.", bold_prefix="Immediate Focus: ")

    add_callout(
        "While the original EU Project Description of Action (Page 5) referenced EAV Bigogwe, TSS Kisaro, and EAV Rushashi as provisional targets, the technical PTU design and 2026 workplan confirm Kisaro TSS and EFA Nyagahanga TSS as the primary CoVE investment sites.",
        title="Institutional Site Context",
        alert_type="note"
    )

    add_heading_2("5.2. Thematic Workstream Status")
    add_bullet("CBT/CBA Curricula Modernization: National-level alignment with RTB and NESA curriculum frameworks. Transitioning traditional syllabus towards practical, green, and competence-based standards.", bold_prefix="Curricula Reform: ")
    add_bullet("Practical Training Units (PTU): Establishing the dual model of student practical learning and self-sustaining production enterprise across target schools.", bold_prefix="PTU Ecosystem: ")
    add_bullet("Recognition of Prior Learning (RPL): Building technical assistance roadmap for RTB to assess and certify informal agricultural skills (Activity A.1.2.9; Oct 2026–2028).", bold_prefix="RPL in Agriculture: ")
    add_bullet("Private Sector Linkages & Workplace Learning: Engaging agricultural enterprises for teacher industrial attachments (Activity A.1.3.6; 2027) and student work-based learning.", bold_prefix="Industry Linkages: ")

    # ==========================================
    # 6. RISKS, BOTTLENECKS & DEPENDENCIES
    # ==========================================
    add_heading_1("6. Risks, Bottlenecks & Critical Dependencies")
    
    add_paragraph("A comprehensive analysis of the 2026 TVET activities reveals five critical risks, bottlenecks, and inter-activity dependencies that require proactive management:")

    # Risks Table
    risk_headers = ["Risk / Bottleneck Area", "Affected Activities", "Risk Severity", "Operational Impact", "Mitigation & Action Required"]
    risk_data = [
        [
            "Dependency on Aimable's Needs Assessment Study",
            "A.1.2.6 (Modular Short Courses & 5 sub-activities)",
            "HIGH",
            "Contracting 3 consultants (1.2.6.1), content validation (1.2.6.2), and subsequent training of teachers (1.2.6.3), farmers (1.2.6.4), and students (1.2.6.5) are blocked.",
            "Meet immediately with Aimable/MEAL Lead to obtain preliminary needs assessment findings and define short-course topic priority list to launch TOR drafting without further delay."
        ],
        [
            "Accreditation Framework Overlap & Duplication Risk",
            "A.1.2.2 (Accreditation Framework) & A.1.3.1 (Excellency Committee)",
            "HIGH",
            "Other stakeholders (ETF/RTB/NESA) are currently developing CoE self-assessment and accreditation tools. High risk of duplicated effort and uncoordinated standards.",
            "Hold strategic consultation with RTB, NESA, and ETF. Assess whether to adopt/adapt national framework rather than procuring independent study; align A.1.3.1 timeline accordingly."
        ],
        [
            "Teacher In-Company Training Dependency on CPD Plan",
            "A.1.3.6 (Teacher On-the-Job Training & 3 sub-activities)",
            "MEDIUM",
            "Selection of 4 teacher batches for enterprise attachment depends on teacher participation and evaluation under this year's Continuous Professional Development (CPD) plan.",
            "Coordinate with RTB Teacher Training Unit to monitor CPD rollout; finalize private sector enterprise mapping and host company criteria in Q4 2026."
        ],
        [
            "Resource Centre Setup Deferral to 2028",
            "A.1.3.8 (Resource Centre Setup & Equipping)",
            "MEDIUM",
            "Initially scheduled for October 2026, the activity has been postponed to 2028 pending structural infrastructure readiness and confirmed space allocation at target CoVEs.",
            "Maintain ongoing dialogue with school principals at Kisaro and Nyagahanga regarding space allocation; ensure ICT/library procurement readiness for 2028."
        ],
        [
            "Consultant Congestion & Validation Bottleneck",
            "A.1.2.3, A.1.2.4, A.1.2.5, A.1.2.7 (Validation Workshops)",
            "MEDIUM",
            "Simultaneous delivery of draft guidelines, digital modules, GEDI review, and trainer's guide creates logistical and technical review congestion for RTB and EF.",
            "Establish a sequenced validation workshop schedule across Q3–Q4 2026; group related workshops where feasible to optimize stakeholder attendance and budget."
        ]
    ]
    add_styled_table(risk_headers, risk_data, col_widths=[1.3, 1.0, 0.8, 1.5, 1.67])

    # ==========================================
    # 7. DECISIONS / APPROVALS REQUIRED
    # ==========================================
    add_heading_1("7. Decisions & Approvals Required")
    
    add_paragraph("The incoming TVET Expert and Project Team Leader must address five strategic decisions to unlock delayed streams and maintain implementation momentum:")

    add_heading_2("Decision 1: Strategic Positioning on Agri CoE Accreditation Framework (Activity A.1.2.2)")
    add_paragraph("The project workplan notes: ", bold_prefix="Context & Problem: ")
    p_d1 = doc.paragraphs[-1]
    p_d1.add_run("'Discuss the feasibility as this is being developed by other stakeholders.'")
    p_d1.add_run(" RTB and the European Training Foundation (ETF) have advanced national guidelines for TVET Centres of Excellence.")
    add_bullet("Option A (Recommended): Partner with RTB/ETF to adapt the existing International Self-Assessment Tool for Centres of Vocational Excellence to the Rwandan agricultural context, focusing EF technical assistance on integrating climate action, gender, and disability metrics.", bold_prefix="Proposed Action: ")
    add_bullet("Option B: Contract an independent consultancy to develop a separate accreditation framework (Risk of misalignment with RTB national standards).", bold_prefix="Alternative: ")
    add_bullet("Decision Maker: Project Team Leader in consultation with RTB Senior Management. Deadline: 15 September 2026.", bold_prefix="Decision Authority: ")

    add_heading_2("Decision 2: Fast-Tracking Needs Assessment Data for Modular Short Courses (Activity A.1.2.6)")
    add_paragraph("Activity A.1.2.6 is scheduled to start in September 2026 but is currently stalled awaiting finalization of the Needs Assessment study led by Aimable.", bold_prefix="Context & Problem: ")
    add_bullet("Request the MEAL/Assessment team to release an interim technical briefing on priority skill gaps in sustainable agriculture, circular economy, and nutrition by 10 September 2026 to enable immediate drafting of TORs for 3 module development consultants.", bold_prefix="Proposed Action: ")
    add_bullet("Decision Maker: Project Team Leader & MEAL Lead. Deadline: 10 September 2026.", bold_prefix="Decision Authority: ")

    add_heading_2("Decision 3: Formal Validation of Reviewed PTU Technical Design Blueprint")
    add_paragraph("The technical review of the Practical Training Unit (PTU) Model Design for Kisaro TSS and EFA Nyagahanga TSS is complete and documented in ", bold_prefix="Context & Problem: ")
    p_d3 = doc.paragraphs[-1]
    p_d3.add_run("MSGS_first_draft_PTU_Model_Design_1_REVIEWED.docx")
    p_d3.add_run(". Formal project approval is required to initiate site preparation and infrastructure procurement.")
    add_bullet("Organize a technical validation session with the Project Infrastructure/Procurement team, RTB, and school leadership to endorse the spatial zoning, machinery lists, and dual governance structure.", bold_prefix="Proposed Action: ")
    add_bullet("Decision Maker: Project Team Leader & Expertise France Management. Deadline: 30 September 2026.", bold_prefix="Decision Authority: ")

    add_heading_2("Decision 4: Procurement Launch for CoE Governance Model Studies (Activity A.1.2.1)")
    add_paragraph("Activity A.1.2.1 (Develop a governance model for Agri CoEs, including financial sustainability mechanisms) is planned for September 2026–2027.", bold_prefix="Context & Problem: ")
    add_bullet("Approve the Terms of Reference (TOR) for contracting expertise to conduct Phase 1 (situational diagnostic) and Phase 2 (governance model formulation) adapted to Kisaro TSS and EFA Nyagahanga TSS.", bold_prefix="Proposed Action: ")
    add_bullet("Decision Maker: TVET Expert & Procurement Officer. Deadline: 20 September 2026.", bold_prefix="Decision Authority: ")

    add_heading_2("Decision 5: Endorsement of Inter-Agency Technical Assistance Roadmap for RPL (Activity A.1.2.9)")
    add_paragraph("Activity A.1.2.9 supports RTB with technical assistance to initiate Recognition of Prior Learning (RPL) in Agriculture, scheduled from October 2026 through 2028.", bold_prefix="Context & Problem: ")
    add_bullet("Convene a high-level technical meeting with RTB and RAB to formally agree on the scope of project support (workshop participation, assessor training curriculum, and pilot certification).", bold_prefix="Proposed Action: ")
    add_bullet("Decision Maker: Project Team Leader & RTB Director General. Deadline: 15 October 2026.", bold_prefix="Decision Authority: ")

    # ==========================================
    # 8. IMMEDIATE PRIORITIES FOR SUCCESSOR
    # ==========================================
    add_heading_1("8. Immediate Priorities for the Successor (30-60-90 Day Plan)")
    
    add_paragraph("To maintain project momentum and ensure high-quality delivery, the incoming TVET Expert should execute the following prioritized roadmap:")

    add_heading_2("Phase 1: Days 1–30 (Immediate Onboarding & In-Flight Control)")
    add_bullet("Consultant Contract & Deliverable Review: Assume immediate technical supervision of the 4 active consultancies: (1) Assessment Guidelines (A.1.2.3), (2) Curricula Digital Content (A.1.2.4), (3) Gender Review Modules (A.1.2.5), and (4) Trainer's Guide (A.1.2.7). Review current drafts against TOR quality benchmarks.", bold_prefix="Day 1–10: ")
    add_bullet("Needs Assessment Alignment: Meet with Aimable and the MEAL team to unblock the topic selection for Modular Short Courses (A.1.2.6); extract key findings for circular economy and green agriculture training modules.", bold_prefix="Day 10–15: ")
    add_bullet("Accreditation Strategy Briefing: Hold strategic alignment session with the Project Team Leader, RTB, and ETF focal points regarding the Agri CoE Accreditation Framework (A.1.2.2).", bold_prefix="Day 15–20: ")
    add_bullet("School Leadership Engagement: Conduct introductory visits or virtual meetings with the Principals and PTU managers of Kisaro TSS and EFA Nyagahanga TSS to review PTU technical design expectations.", bold_prefix="Day 20–30: ")

    add_heading_2("Phase 2: Days 31–60 (Validation & Procurement Launch)")
    add_bullet("National Validation Workshops: Organize and facilitate stakeholder validation workshops for the completed Assessment Guidelines (1.2.3.3) and Gender-Responsive Trainer's Guide (1.2.7.2).", bold_prefix="Day 31–45: ")
    add_bullet("Modular Short Courses Procurement: Draft TORs and initiate procurement for 3 consultants to develop training contents for digital modular short courses (1.2.6.1).", bold_prefix="Day 40–50: ")
    add_bullet("Governance Model TOR Launch: Finalize and advertise the procurement notice for the CoE Governance Model study (1.2.1.1) in accordance with the September 2026 workplan schedule.", bold_prefix="Day 45–55: ")
    add_bullet("CBT/CBA Capacity Building Consultation: Convene the consultative meeting with CBT/CBA and GEDI stakeholders (1.2.8.1) to define training modules for lead trainers.", bold_prefix="Day 50–60: ")

    add_heading_2("Phase 3: Days 61–90 (Implementation & Strategic Rollout)")
    add_bullet("Curricula Digital Content Validation: Conduct national validation workshops for the finalized digital learning modules and CBT/CBA manuals across 3 trades (1.2.4.2 & 1.2.5.2).", bold_prefix="Day 61–75: ")
    add_bullet("RPL Technical Assistance Mobilization: Finalize the technical assistance agreement with RTB for RPL in Agriculture (1.2.9.1) and participate in initial stakeholder planning events.", bold_prefix="Day 70–80: ")
    add_bullet("Training of Trainers (ToT) Execution: Launch the first ToT session for selected Lead Trainers on CBT/CBA and GEDI methodology (1.2.8.2).", bold_prefix="Day 75–90: ")
    add_bullet("PTU Implementation Oversight: Coordinate with the infrastructure team on the ground preparation and equipment procurement pipeline for Kisaro and Nyagahanga PTUs.", bold_prefix="Day 80–90: ")

    # ==========================================
    # 9. HANDOVER ACTION TRACKER
    # ==========================================
    add_heading_1("9. Handover Action Tracker")
    
    add_paragraph("The following action tracker provides an exhaustive, activity-by-activity status record for all 26 workplan lines and core project issues. To ensure complete documentation integrity, unconfirmed parameters are explicitly designated as 'To be confirmed'.")

    tracker_headers = ["Activity / Issue", "Status", "Completed Work", "Outstanding Work", "Next Action", "Responsible Person", "Deadline", "Priority"]
    
    tracker_data = [
        [
            "A.1.2.1 Develop CoE Governance Model & Financial Sustainability",
            "PLANNED",
            "Baseline terms outlined in project document.",
            "Phase 1 & 2 diagnostic studies; 2 adapted governance models.",
            "Finalize TOR; launch procurement for governance expertise.",
            "TVET Expert / Procurement",
            "Sept 2026–2027",
            "HIGH"
        ],
        [
            "1.2.1.1 Contract expertise for governance studies (Phase 1 & 2)",
            "PLANNED",
            "Draft scope of work prepared.",
            "Procurement, contracting, and study execution.",
            "Publish tender; evaluate technical proposals.",
            "Procurement / TVET Expert",
            "September 2026",
            "HIGH"
        ],
        [
            "1.2.1.2 Develop governance models for 2 CoEs (Kisaro & Nyagahanga)",
            "PLANNED",
            "Initial institutional mapping.",
            "Drafting and adapting governance models for 2 schools.",
            "Supervise consultant field data collection at 2 CoEs.",
            "Consultant / TVET Expert",
            "To be confirmed",
            "MEDIUM"
        ],
        [
            "A.1.2.2 Develop Accreditation Framework for Agri CoEs",
            "PLANNED",
            "Initial review of ETF CoVE assessment tools.",
            "Benchmarking study, contextual adaptation, validation.",
            "Convene alignment meeting with RTB/ETF on tool duplication.",
            "TVET Expert / Team Leader",
            "Sept 2026–2027",
            "HIGH"
        ],
        [
            "1.2.2.1 Conduct benchmarking study of agri CoE accreditation",
            "PLANNED",
            "None to date.",
            "Comparative analysis of regional/international frameworks.",
            "Determine feasibility based on external stakeholder progress.",
            "TVET Expert / Consultant",
            "To be confirmed",
            "MEDIUM"
        ],
        [
            "1.2.2.2 Adapt best practices to Rwanda & integrate GEDI/climate",
            "PLANNED",
            "None to date.",
            "Integration of climate, gender, and inclusion metrics.",
            "Draft criteria matrix aligned with Rwandan TVET law.",
            "TVET Expert / Consultant",
            "To be confirmed",
            "MEDIUM"
        ],
        [
            "1.2.2.3 Co-organize validation workshop for accreditation framework",
            "PLANNED",
            "None to date.",
            "Multi-stakeholder validation event.",
            "Plan workshop once draft framework is finalized.",
            "TVET Expert / RTB",
            "To be confirmed",
            "MEDIUM"
        ],
        [
            "A.1.2.3 Develop Assessment Guidelines for Agri CoEs",
            "IN PROGRESS",
            "Consultant contracted; draft guidelines drafted.",
            "Final draft refinement; validation workshop.",
            "Review draft guidelines; organize validation workshop.",
            "TVET Expert / Consultant",
            "Q3 2026 (TBC)",
            "HIGH"
        ],
        [
            "1.2.3.1 Contract consultant for assessment guidelines (experiential)",
            "IN PROGRESS",
            "Consultant contracted and engaged.",
            "Deliverable submission and administrative closure.",
            "Monitor deliverable milestones against contract schedule.",
            "TVET Expert / Consultant",
            "To be confirmed",
            "HIGH"
        ],
        [
            "1.2.3.2 Develop assessment guidelines (experiential learning model)",
            "IN PROGRESS",
            "First draft submitted for review.",
            "Integration of technical feedback from TVET Expert.",
            "Finalize consolidated draft guideline document.",
            "Consultant",
            "To be confirmed",
            "HIGH"
        ],
        [
            "1.2.3.3 Organize workshop to validate assessment guidelines",
            "PLANNED",
            "Workshop concept outlined.",
            "Participant invitations, venue logistics, workshop execution.",
            "Draft workshop concept note using standardized template.",
            "TVET Expert / RTB",
            "To be confirmed",
            "HIGH"
        ],
        [
            "A.1.2.4 Develop/Review Curricula in Agri, Animal Health, Food Proc.",
            "IN PROGRESS",
            "Content outlines drafted for 3 trades.",
            "Finalization of 2 digital module sets; 2 validation workshops.",
            "Conduct technical review of digital training manuals.",
            "TVET Expert / RTB",
            "Q3–Q4 2026",
            "HIGH"
        ],
        [
            "1.2.4.1 Develop Digital Learning Content & Training Manuals (Target: 2)",
            "IN PROGRESS",
            "Draft digital modules under preparation.",
            "Complete draft modules for 3 trades.",
            "Review modules for CBT/CBA and green skill compliance.",
            "Consultants / TVET Expert",
            "To be confirmed",
            "HIGH"
        ],
        [
            "1.2.4.2 Organize workshops to validate developed modules (Target: 2)",
            "PLANNED",
            "None to date.",
            "2 stakeholder validation workshops.",
            "Schedule validation workshops upon receipt of final drafts.",
            "TVET Expert / RTB",
            "To be confirmed",
            "HIGH"
        ],
        [
            "A.1.2.5 Review Curricula for Gender-Responsiveness and Inclusion",
            "IN PROGRESS",
            "Gender screening criteria drafted.",
            "Complete trade reviews; draft 2 modules; validation workshop.",
            "Coordinate with Gender Expert on draft trade modules.",
            "TVET Expert / Gender Exp",
            "Q3–Q4 2026",
            "HIGH"
        ],
        [
            "1.2.5.1 Conduct curricula review & develop GEDI modules (Target: 2)",
            "IN PROGRESS",
            "Trade review underway.",
            "Finalization of 2 digital GEDI modules.",
            "Review module content with gender advocacy focal points.",
            "Consultant / Gender Exp",
            "To be confirmed",
            "HIGH"
        ],
        [
            "1.2.5.2 Organize workshops to validate new GEDI modules (Target: 2)",
            "PLANNED",
            "None to date.",
            "2 validation workshops.",
            "Plan joint validation workshop with A.1.2.4 where feasible.",
            "TVET Expert / Gender Exp",
            "To be confirmed",
            "MEDIUM"
        ],
        [
            "A.1.2.6 Develop Modular Short Courses (Circular Econ, Gender, Nutr.)",
            "DELAYED / AT RISK",
            "Activity scope defined in results framework.",
            "3 consultant contracts; 3 contents; teacher/farmer/student training.",
            "Obtain Needs Assessment data from Aimable; draft TORs.",
            "TVET Expert / MEAL Lead",
            "From Sept 2026",
            "HIGH"
        ],
        [
            "1.2.6.1 Contract 3 consultants for short course topics (Digital)",
            "PLANNED",
            "None to date (delayed by Needs Assessment).",
            "TOR formulation, tendering, contracting 3 consultants.",
            "Finalize TORs based on preliminary assessment findings.",
            "TVET Expert / Procurement",
            "September 2026",
            "HIGH"
        ],
        [
            "1.2.6.2 Validate developed short course content via workshop (Target: 3)",
            "PLANNED",
            "None to date.",
            "3 content validation workshops.",
            "Schedule validation sessions post-content drafting.",
            "TVET Expert / RTB",
            "To be confirmed",
            "MEDIUM"
        ],
        [
            "1.2.6.3 Conduct training to trainers/teachers (Target: 3 sessions)",
            "PLANNED",
            "None to date.",
            "3 teacher training sessions delivered.",
            "Select trainee teachers in coordination with CoVE principals.",
            "TVET Expert / Master Train",
            "To be confirmed",
            "MEDIUM"
        ],
        [
            "1.2.6.4 Conduct training to farmers",
            "PLANNED",
            "None to date.",
            "Farmer outreach training sessions delivered.",
            "Coordinate with RAB/FFS networks surrounding CoVEs.",
            "TVET Expert / RAB",
            "To be confirmed",
            "MEDIUM"
        ],
        [
            "1.2.6.5 Conduct training to students",
            "PLANNED",
            "None to date.",
            "Student short-course modules delivered.",
            "Embed short courses into target CoVE academic calendar.",
            "TVET Expert / CoVE Staff",
            "To be confirmed",
            "MEDIUM"
        ],
        [
            "A.1.2.7 Develop Gender-Responsive Trainers' Guide",
            "IN PROGRESS",
            "Consultant contracted; draft guide drafted.",
            "Final draft refinement; validation workshop.",
            "Review draft trainer's guide; organize validation workshop.",
            "TVET Expert / Consultant",
            "Q3 2026 (TBC)",
            "HIGH"
        ],
        [
            "1.2.7.1 Contract consultant & develop trainer's guide",
            "IN PROGRESS",
            "Consultant contracted; initial draft developed.",
            "Finalization of complete trainer's guide document.",
            "Perform comprehensive technical QA on draft guide.",
            "TVET Expert / Consultant",
            "To be confirmed",
            "HIGH"
        ],
        [
            "1.2.7.2 Conduct workshop to validate trainer's guide",
            "PLANNED",
            "None to date.",
            "National validation workshop.",
            "Draft workshop concept note and invite lead trainers.",
            "TVET Expert / RTB",
            "To be confirmed",
            "HIGH"
        ],
        [
            "A.1.2.8 CBT/CBA Capacity Building for Training/Eval Framework Staff",
            "PLANNED",
            "Concept framework established in concept note template.",
            "1 consultative meeting; 1 Lead ToT; 1 Trainer ToT (4 sessions).",
            "Convene consultative meeting with RTB & GEDI stakeholders.",
            "TVET Expert / RTB",
            "Aug 2026–2027",
            "HIGH"
        ],
        [
            "1.2.8.1 Consultative meeting on CBT/CBA & GEDI focus (Target: 1)",
            "PLANNED",
            "Agenda points drafted.",
            "Consultative meeting conducted with stakeholders.",
            "Issue meeting invitations to RTB, NESA, and GEDI partners.",
            "TVET Expert / RTB",
            "August/Sept 2026",
            "HIGH"
        ],
        [
            "1.2.8.2 Conduct Training of Trainers (ToT) to Lead Trainers (Target: 1)",
            "PLANNED",
            "Training modules outlined.",
            "Delivery of Lead ToT session.",
            "Finalize participant list of Lead Trainers across districts.",
            "TVET Expert / Master Train",
            "To be confirmed",
            "HIGH"
        ],
        [
            "1.2.8.3 Conduct Training of Trainers to selected trainers (Target: 1)",
            "PLANNED",
            "None to date.",
            "Delivery of general Trainer ToT session.",
            "Schedule training following completion of Lead ToT.",
            "TVET Expert / Lead Train",
            "To be confirmed",
            "MEDIUM"
        ],
        [
            "A.1.2.9 Support Recognition of Prior Learning (RPL) in Agriculture",
            "PLANNED",
            "Preliminary discussions with RTB.",
            "TA to plan RPL; 1 stakeholder workshop; 1 assessor training.",
            "Formalize technical assistance roadmap with RTB RPL unit.",
            "TVET Expert / RTB",
            "Oct 2026–2028",
            "HIGH"
        ],
        [
            "1.2.9.1 Support RTB with TA to plan/start RPL process in Agriculture",
            "PLANNED",
            "Scope of TA identified.",
            "Provision of ongoing technical advisory to RTB.",
            "Draft TA terms of reference and schedule joint kickoff.",
            "TVET Expert / RTB",
            "October 2026",
            "HIGH"
        ],
        [
            "1.2.9.2 Co-participate in workshops/events for RPL development (Target: 1)",
            "PLANNED",
            "None to date.",
            "1 stakeholder workshop co-facilitated.",
            "Coordinate date and project contribution with RTB.",
            "TVET Expert / RTB",
            "To be confirmed",
            "MEDIUM"
        ],
        [
            "1.2.9.3 Co-participate in training for RPL assessors/evaluators (Target: 1)",
            "PLANNED",
            "None to date.",
            "1 assessor training conducted.",
            "Review RPL assessor training curriculum for agriculture.",
            "TVET Expert / RTB",
            "To be confirmed",
            "MEDIUM"
        ],
        [
            "A.1.3.1 Set up & build capacity of Excellency Committees (Target: 2)",
            "PLANNED",
            "Committees mapped to 2 CoVE sites.",
            "Establishment of 2 committees; 2 Excellency Action Plans.",
            "Align committee terms with finalized accreditation criteria.",
            "TVET Expert / CoVEs",
            "2027 (Next Year)",
            "MEDIUM"
        ],
        [
            "A.1.3.6 Support On-the-Job Training of Teaching Staff in Companies (Target: 4)",
            "PLANNED",
            "Enterprise partnership framework outlined in PTU doc.",
            "Identify host companies; sign agreements; place & visit trainers.",
            "Review teacher CPD response; draft host company agreements.",
            "TVET Expert / Private Sec",
            "2027 (Next Year)",
            "MEDIUM"
        ],
        [
            "1.3.6.1 Identify companies to host designated trainers/teachers",
            "PLANNED",
            "Preliminary private sector mapping.",
            "Final roster of qualified agricultural host enterprises.",
            "Engage PSF Agriculture Chamber and regional agribusinesses.",
            "TVET Expert / PSF",
            "To be confirmed",
            "MEDIUM"
        ],
        [
            "1.3.6.2 Sign agreements with host companies & place trainers",
            "PLANNED",
            "Agreement template outlined.",
            "Signed MoUs and successful placement of 4 teacher batches.",
            "Draft bipartite placement agreements with host firms.",
            "TVET Expert / EF Legal",
            "To be confirmed",
            "MEDIUM"
        ],
        [
            "1.3.6.3 Visit & monitor trainers under on-the-job training",
            "PLANNED",
            "None to date.",
            "Supervisory field monitoring visits conducted.",
            "Establish supervisory logbook and assessment rubric.",
            "TVET Expert",
            "To be confirmed",
            "MEDIUM"
        ],
        [
            "A.1.3.8 Set up, equip and support management of Resource Centre",
            "PLANNED",
            "Resource centre equipment list outlined in PTU design.",
            "Procurement of materials; physical setup; operational launch.",
            "Monitor school space readiness; prepare procurement for 2028.",
            "TVET Expert / Procurement",
            "Deferred to 2028",
            "LOW"
        ],
        [
            "PTU Model Technical Design & Blueprint Validation",
            "COMPLETED",
            "Full technical review of 50-page PTU design doc.",
            "Architectural translation and site civil works tendering.",
            "Endorse reviewed PTU blueprint for Kisaro & Nyagahanga.",
            "TVET Expert / Team Leader",
            "Completed (Aug 2026)",
            "HIGH"
        ],
        [
            "Capacity Building Concept Note Framework Standardisation",
            "COMPLETED",
            "Standardized Training Concept Note Template produced.",
            "Application of template across all upcoming training events.",
            "Ensure all subsequent workshop concept notes follow template.",
            "TVET Expert / Successor",
            "Completed (Aug 2026)",
            "HIGH"
        ]
    ]
    
    add_styled_table(tracker_headers, tracker_data, col_widths=[1.2, 0.8, 0.9, 0.9, 0.9, 0.7, 0.5, 0.37])

    # ==========================================
    # 10. INFORMATION GAPS & ITEMS TO CONFIRM
    # ==========================================
    add_heading_1("10. Information Gaps, Spreadsheet Inconsistencies & Items to Confirm")
    
    add_paragraph("A rigorous audit of the project's source data files revealed several technical inconsistencies, formula errors, and unconfirmed parameters that the successor must rectify in coordination with the MEAL and Project Management teams:")

    add_heading_2("10.1. Spreadsheet Inconsistencies & Technical Formula Errors")
    add_bullet("Formula Error (#DIV/0!): In Sheet 1 ('2026 Activities'), Row 48 (Activity A.1.3.8 - Resource Centre) displays a '#DIV/0!' formula error in the '% Progress' column. This resulted from an Excel formula attempting to divide achieved progress (0) by an unpopulated numerical target cell. In Sheet 2 ('2026 Activities improved'), the error was cleaned by removing the redundant columns, but the master logframe requires a standardized formula rule (e.g., IFERROR).", bold_prefix="Formula Error in Sheet 1: ")
    add_bullet("Progress Tracking Discrepancy: Sheet 1 logs numerical 'Achieved' values as 0 and '% Progress' as 0% for activities marked as 'IN PROGRESS' (such as A.1.2.3, A.1.2.4, A.1.2.5, A.1.2.7). This creates a misleading impression in automated dashboard rollups that no work has occurred, despite extensive drafting and consultant contracting. The successor should ensure qualitative milestones are properly mapped to milestone percentages (e.g., 50% for draft deliverable submitted).", bold_prefix="Status vs Percentage Mismatch: ")
    add_bullet("Duplicate Worksheets: The workbook contains two overlapping worksheets: '2026 Activities' (9 columns with numerical achievement fields) and '2026 Activities improved' (7 columns cleaned). Sheet 2 should be officially designated as the active working matrix, and legacy sheets should be archived to prevent version confusion.", bold_prefix="Worksheet Versioning: ")

    add_heading_2("10.2. Missing Information & Items Requiring Formal Confirmation")
    add_bullet("Exact Consultant Contract Deadlines: The spreadsheet marks active sub-activities (1.2.3.1, 1.2.3.2, 1.2.4.1, 1.2.5.1, 1.2.7.1) as 'IN PROGRESS' but leaves the calendar completion dates blank. The successor must review the physical contracts with the Procurement Officer to confirm exact deliverable due dates.", bold_prefix="Active Contract Deadlines: ")
    add_bullet("Unspecified Quantitative Targets: Several sub-activities have blank target cells (e.g., 1.2.1.1, 1.2.1.2, 1.2.2.1, 1.2.2.2, 1.2.2.3, 1.2.3.1, 1.2.3.2, 1.2.3.3, 1.2.6.4, 1.2.6.5, 1.2.7.1, 1.2.7.2, 1.2.9.1, 1.3.6.1, 1.3.6.2, 1.3.6.3, A.1.3.8). These must be populated with specific target metrics (e.g., number of reports, number of participants, number of MoUs signed) during the upcoming workplan update.", bold_prefix="Sub-Activity Target Metrics: ")
    add_bullet("Delivery Date of Aimable's Needs Assessment: The exact delivery milestone for the Needs Assessment study is not formally logged in the TVET matrix, directly jeopardizing the start of Activity A.1.2.6. A formal delivery date must be confirmed with the Project Team Leader.", bold_prefix="Needs Assessment Milestone: ")
    add_bullet("Final Target School List Confirmation: The original EU Description of Action mentions EAV Bigogwe, TSS Kisaro, and EAV Rushashi (Page 5), whereas the PTU Technical Design document and current workplan focus specifically on Kisaro TSS (Rulindo) and EFA Nyagahanga TSS (Gatsibo). Formal confirmation of the total scope of schools supported across the 5-year project lifecycle should be re-validated with Project Management.", bold_prefix="School List Reconciliation: ")

    # ==========================================
    # 11. ANNEXES & REFERENCE REPOSITORY
    # ==========================================
    add_heading_1("11. Annex: Key Documents & Reference Repository")
    
    add_paragraph("All primary technical files, templates, reference blueprints, and workplan matrices are organized in the project folder as detailed below:")

    ref_headers = ["Document File Name", "Document Type & Description", "Key Technical Content", "Location / Path"]
    ref_data = [
        [
            "TVET_Expert_2026_Activities.xlsx",
            "Master Workplan & Activity Tracker",
            "Complete log of 26 TVET activities, indicator targets, status tracking, periods, and dependency comments.",
            "Project Root Directory"
        ],
        [
            "MSGS_first_draft_PTU_Model_Design_1_REVIEWED.docx",
            "Technical Design Blueprint (Reviewed)",
            "50+ page detailed engineering, pedagogical, zoning, and governance blueprint for PTUs at Kisaro TSS and EFA Nyagahanga TSS.",
            "Project Root Directory"
        ],
        [
            "Training Concept Note Template.docx",
            "Standardized Pedagogical Template",
            "Standardized template for all capacity building events, training objectives (Bloom's Taxonomy), results framework links, and M&E.",
            "Project Root Directory"
        ],
        [
            "TVET Agri_Document.pdf",
            "EU Description of Action (DoA)",
            "52-page master project contract (No. 700002154) detailing overall objectives, specific outcomes, results logframe, and governance.",
            "Project Root Directory"
        ],
        [
            "ETF_TVET CoVEs Rwanda.pdf",
            "Policy & Benchmark Reference",
            "European Training Foundation policy study on TVET Centres of Excellence in Rwanda (governance, services, self-assessment).",
            "Project Root Directory"
        ]
    ]
    add_styled_table(ref_headers, ref_data, col_widths=[1.5, 1.3, 2.0, 1.47])

    # Final Sign-off block
    add_paragraph("Report Prepared and Submitted by:", bold_prefix="Document Sign-Off: ", space_after=2)
    add_paragraph("Outgoing TVET Expert, TVET Agri – Ubuhinzi Skills+ Project, Expertise France Rwanda", bold_prefix="Author: ")
    add_paragraph("Incoming TVET Expert & Project Team Leader, Expertise France Rwanda", bold_prefix="Received by: ")
    add_paragraph("August 2026 | Kigali, Rwanda", bold_prefix="Date & Location: ")

    output_path = "/workspaces/TVET-Agri/TVET_Agri_TVET_Expert_Handover_Report.docx"
    doc.save(output_path)
    print(f"Report successfully created at {output_path}")

if __name__ == "__main__":
    build_handover_report()
